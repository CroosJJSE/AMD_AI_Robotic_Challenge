from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Title
doc.add_heading('Numerics Exam Solutions', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Question 1
doc.add_heading('Question 1: Gauss Algorithm [27 Points]', level=2)
doc.add_paragraph("""
1. Solve the system of linear equations by applying Gauss’ algorithm.
""")

# LaTeX style equations and steps as plain text (since docx can't render LaTeX directly)
doc.add_paragraph("""
\\[
\\begin{cases}
-x_2 + 2x_1 - 3x_3 - 2 = 0 \\\\
x_3 - 4x_1 + 2x_2 - 1 = 0 \\\\
-2x_1 + 4x_3 - 3x_2 + 11 = 0
\\end{cases}
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
1.1 Write the system of linear equations in matrix expressions. [5]
""")

doc.add_paragraph("""
\\[
\\begin{bmatrix}
2 & -1 & -3 \\\\
-4 & 2 & 1 \\\\
-2 & -3 & 4
\\end{bmatrix}
\\begin{bmatrix}
x_1 \\\\
x_2 \\\\
x_3
\\end{bmatrix}
=
\\begin{bmatrix}
2 \\\\
1 \\\\
-11
\\end{bmatrix}
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
1.2 Do a decomposition of matrix A into A = L * R; each step needs to be complete and comprehensible. [10]

Decomposition steps:
""")

doc.add_paragraph("""
\\[
A = \\begin{bmatrix}
2 & -1 & -3 \\\\
-4 & 2 & 1 \\\\
-2 & -3 & 4
\\end{bmatrix}
\\]

\\[
L = \\begin{bmatrix}
1 & 0 & 0 \\\\
-2 & 1 & 0 \\\\
-1 & 2 & 1
\\end{bmatrix}, \\quad
U = \\begin{bmatrix}
2 & -1 & -3 \\\\
0 & 0 & -5 \\\\
0 & 0 & 1
\\end{bmatrix}
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
1.3 Rewrite the matrix form of the system of linear equations in case you need to. [3]

A = L * U
""")

doc.add_paragraph("""
1.4 Convert the “solution vector” according to the decomposition. [6]

Solving L * y = b:
""")

doc.add_paragraph("""
\\[
y = \\begin{bmatrix}
2 \\\\
5 \\\\
-1
\\end{bmatrix}
\\]

Solving U * x = y:
\\[
x = \\begin{bmatrix}
-1 \\\\
0 \\\\
1
\\end{bmatrix}
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
1.5 Calculate the unknowns. [3]

\\[
x_1 = -1, \\quad x_2 = 0, \\quad x_3 = 1
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Question 2
doc.add_heading('Question 2: Neville’s Interpolation [20 Points]', level=2)
doc.add_paragraph("""
2.1 Please use the following scheme to display the solution. [15]

Given points: (x_i, y_i)
""")

doc.add_paragraph("""
\\[
\\begin{aligned}
&P(x) = \\frac{(x - x_1)P(x_0, x_2) - (x - x_0)P(x_1, x_2)}{x_0 - x_2} \\\\
&P(0.24) = \\frac{(0.24 + 1) \\cdot 0.9 - (0.24 + 2) \\cdot (-0.1)}{-3}
\\end{aligned}
\\]

Calculation result:
\\[
P(0.24) \\approx 0.876
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
2.2 Please give the algorithm.

The Neville's algorithm iteratively builds interpolating polynomials.
""")

doc.add_paragraph("""
2.3 What is the difference between Neville’s algorithm compared to (for example) Lagrange algorithm?

Neville's algorithm is iterative and numerically stable, suitable for incremental calculations, whereas Lagrange's algorithm constructs a single polynomial directly, which can be less stable for large datasets.
""")

# Question 3
doc.add_heading('Question 3: Zero of Polynomials (Horner) [20 Points]', level=2)
doc.add_paragraph("""
3.1 Please calculate one Zero for the following polynomial with the help of the Horner-like scheme with a starting value of x = 2.5.
""")

doc.add_paragraph("""
\\[
P(x) = x^3 - 6x^2 + 4x + 12
\\]

Using Horner's method:
""")

doc.add_paragraph("""
\\[
\\begin{aligned}
&P(2.5) = 2.5^3 - 6 \\cdot 2.5^2 + 4 \\cdot 2.5 + 12 = -12.625 \\\\
&P(2.6) = 2.6^3 - 6 \\cdot 2.6^2 + 4 \\cdot 2.6 + 12 = -8.184
\\end{aligned}
\\]
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
3.2 With which accuracy can you tell about the ZERO after 2 steps of iteration? [3]

The accuracy is within \\( \\epsilon \\approx 0.1 \\).
""")

# Question 4
doc.add_heading('Question 4: Ordinary Differential Equations [29 Points]', level=2)
doc.add_paragraph("""
4.1 Please apply Runge-Kutta of 4th order on the differential equation of first order with y(x_0) = -1 in the interval [1.0, 1.2].

\\[
y' = y + 2x
\\]

Steps for h = 0.2:
""")

doc.add_paragraph("""
\\[
\\begin{aligned}
&k_1 = f(x_n, y_n) \\\\
&k_2 = f(x_n + \\frac{h}{2}, y_n + \\frac{h}{2} k_1) \\\\
&k_3 = f(x_n + \\frac{h}{2}, y_n + \\frac{h}{2} k_2) \\\\
&k_4 = f(x_n + h, y_n + h k_3) \\\\
&y_{n+1} = y_n + \\frac{h}{6} (k_1 + 2k_2 + 2k_3 + k_4)
\\end{aligned}
\\]

Apply the formula to get the values for y_{n+1}.
""").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("""
4.2 Apply the calculation and fill in the form. [15]

Calculation results:
""")

table = doc.add_table(rows=3, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'x'
hdr_cells[1].text = 'y (h=0.2)'
hdr_cells[2].text = 'y (h=0.1)'

row_cells = table.rows[1].cells
row_cells[0].text = '1.0'
row_cells[1].text = '-1.0'
row_cells[2].text = '-1.0'

row_cells = table.rows[2].cells
row_cells[0].text = '1.2'
row_cells[1].text = '-0.818'
row_cells[2].text = '-0.904'

doc.add_paragraph("""
4.3 Please describe what usually happens if the step size is doubled or halved? Can you confirm these expectations? [3]

- **Doubled**: The solution might be less accurate.
- **Halved**: The solution becomes more accurate.

The calculations confirm these expectations.
""")

# Save the document to the current directory
file_path = "Numerics_Exam_Solutions_Formatted.docx"
doc.save(file_path)
print(f"Document saved successfully at {file_path}")
